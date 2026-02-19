"""Create sample documents for all supported ingestion formats.

Generates: .txt, .pdf, .docx, .eml, .png (for OCR testing).
Run:  python tests/create_sample_docs.py
"""

import sys
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

SAMPLE_DIR = Path(__file__).parent.parent / "data" / "sample_corpus"


# ── 1. Text file ─────────────────────────────────────────────────────
def create_sample_txt():
    """Create a sample .txt memo."""
    path = SAMPLE_DIR / "memo_safety_2024.txt"
    path.write_text(
        "INTERNAL MEMO\n"
        "Date: 2024-03-15\n"
        "From: Sarah Johnson, VP of Operations\n"
        "To: All Department Heads\n\n"
        "Subject: Q1 2024 Compliance Review\n\n"
        "Please be advised that the quarterly compliance review has identified\n"
        "several areas requiring immediate attention:\n\n"
        "1. Environmental Reporting: Three facilities failed to submit monthly\n"
        "   emissions data on time. This is a violation of EPA Form R requirements.\n\n"
        "2. Worker Safety: Incident reports from the Baytown plant show a 15%\n"
        "   increase in reportable injuries. OSHA inspections are expected.\n\n"
        "3. Financial Controls: The internal audit team discovered unauthorized\n"
        "   vendor payments totaling $45,000 in the Houston office.\n\n"
        "All department heads must submit corrective action plans by March 30.\n",
        encoding="utf-8",
    )
    print(f"  Created: {path.name}")


# ── 2. PDF file ──────────────────────────────────────────────────────
def create_sample_pdf():
    """Create sample PDF files."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        print("  Installing reportlab for PDF generation...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

    # PDF 1: Compliance report
    path = SAMPLE_DIR / "compliance_report_2024.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "Annual Compliance Report 2024")
    c.setFont("Helvetica", 11)
    y = 690
    lines = [
        "Prepared by: Legal & Compliance Department",
        "",
        "Executive Summary:",
        "This report covers regulatory compliance across all divisions.",
        "Key findings include worker safety violations in the welding",
        "department and environmental reporting gaps at three facilities.",
        "",
        "Key Finding #1: Welding Department Safety",
        "The Baytown welding facility had 12 reported incidents in Q1 2024,",
        "a 15% increase over the prior quarter. Two incidents involved flash",
        "fires due to improperly maintained fire suppression equipment.",
        "",
        "Key Finding #2: Environmental Compliance",
        "Three facilities failed to file monthly EPA emissions reports.",
        "This puts the company at risk of fines under the Clean Air Act.",
        "",
        "Recommendation: Immediate corrective action required for OSHA",
        "compliance at Baytown. Environmental reporting procedures need",
        "to be automated to prevent future gaps.",
    ]
    for line in lines:
        c.drawString(72, y, line)
        y -= 15
    c.save()
    print(f"  Created: {path.name}")


# ── 3. Word (.docx) file ────────────────────────────────────────────
def create_sample_docx():
    """Create a sample .docx contract amendment."""
    try:
        from docx import Document
    except ImportError:
        print("  SKIP: .docx (pip install python-docx)")
        return

    path = SAMPLE_DIR / "contract_amendment_03.docx"
    doc = Document()
    doc.add_heading("Contract Amendment #3", level=1)
    doc.add_paragraph(
        "This amendment modifies the Master Services Agreement dated January 10, 2023 "
        "between BigThorium Energy Corp ('Company') and Pacific Welding Services ('Contractor')."
    )
    doc.add_heading("Section 1: Scope Changes", level=2)
    doc.add_paragraph(
        "The Contractor shall provide additional welding inspection services at the "
        "Baytown facility. All welders must hold current AWS D1.1 certifications."
    )
    doc.add_heading("Section 2: Compensation", level=2)

    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Service", "Rate ($/hr)", "Estimated Hours"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("Welding Inspection", "85", "200"),
        ("Safety Audit", "120", "40"),
        ("Documentation Review", "95", "80"),
    ]
    for row_idx, (svc, rate, hrs) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = svc
        table.rows[row_idx].cells[1].text = rate
        table.rows[row_idx].cells[2].text = hrs

    doc.add_paragraph("")
    doc.add_paragraph("Effective Date: April 1, 2024")
    doc.add_paragraph("Signed: William Davis, Director of Operations")
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── 4. Email (.eml) file ────────────────────────────────────────────
def create_sample_eml():
    """Create sample .eml email files."""
    # Email 1: Safety incident
    path1 = SAMPLE_DIR / "safety_incident_baytown.eml"
    msg1 = MIMEMultipart()
    msg1["From"] = "Francis Ham <fham@bigthorium.com>"
    msg1["To"] = "William Davis <wdavis@bigthorium.com>"
    msg1["Cc"] = "Arvind Patel <apatel@bigthorium.com>"
    msg1["Date"] = "Thu, 18 Apr 2024 09:32:00 -0500"
    msg1["Subject"] = "URGENT: Safety Incident at Baytown Welding Shop"

    body1 = (
        "Bill,\n\n"
        "I need to bring to your attention a serious safety incident that occurred\n"
        "yesterday at the Baytown welding facility.\n\n"
        "At approximately 2:15 PM, a flash fire broke out in Welding Bay 3 during\n"
        "a routine pipe fitting operation. Two workers sustained minor burns and\n"
        "were transported to Memorial Hermann for evaluation.\n\n"
        "Preliminary investigation suggests:\n"
        "- The fire suppression system in Bay 3 was last inspected 18 months ago\n"
        "  (should be every 6 months per OSHA 1910.252)\n"
        "- The welding curtains were not properly positioned\n"
        "- One worker was not wearing required PPE (face shield)\n\n"
        "I've already contacted our insurance carrier and OSHA regional office.\n"
        "We need to schedule an emergency safety review ASAP.\n\n"
        "Francis Ham\n"
        "Plant Manager, Baytown Operations\n"
    )
    msg1.attach(MIMEText(body1, "plain"))

    # Simulate an attachment
    att = MIMEBase("application", "octet-stream")
    att.set_payload(b"Incident photo placeholder content")
    encoders.encode_base64(att)
    att.add_header("Content-Disposition", "attachment", filename="incident_bay3_photo.jpg")
    msg1.attach(att)
    path1.write_bytes(msg1.as_bytes())
    print(f"  Created: {path1.name}")

    # Email 2: Recruitment discussion
    path2 = SAMPLE_DIR / "recruitment_india.eml"
    msg2 = MIMEMultipart()
    msg2["From"] = "William Davis <wdavis@bigthorium.com>"
    msg2["To"] = "Samantha Jones <sjones@bigthorium.com>"
    msg2["Date"] = "Mon, 16 Jan 2023 14:05:00 -0500"
    msg2["Subject"] = "Re: Skilled Welder Recruitment - India"

    body2 = (
        "Sam,\n\n"
        "Following up on our discussion about recruiting skilled welders from India.\n\n"
        "We need at least 25 certified welders and fitters by Q2. I've been in touch\n"
        "with our recruitment agency in Mumbai. They have several qualified candidates\n"
        "with AWS and ASME certifications.\n\n"
        "We should offer the green card sponsorship package we discussed - it's been\n"
        "effective in attracting top talent. Housing arrangements at the Baytown\n"
        "complex need to be finalized before they arrive.\n\n"
        "Can you coordinate with HR on the visa paperwork?\n\n"
        "Thanks,\n"
        "William Davis\n"
        "Director of Operations\n"
    )
    msg2.attach(MIMEText(body2, "plain"))
    path2.write_bytes(msg2.as_bytes())
    print(f"  Created: {path2.name}")


# ── 5. Image with text (for OCR) ────────────────────────────────────
def create_sample_image():
    """Create a sample image with text for OCR testing."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("  SKIP: .png (pip install Pillow)")
        return

    path = SAMPLE_DIR / "scanned_safety_notice.png"
    img = Image.new("RGB", (900, 500), color="white")
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arial.ttf", 24)
        font_body = ImageFont.truetype("arial.ttf", 16)
    except (OSError, IOError):
        font_title = ImageFont.load_default()
        font_body = font_title

    lines = [
        ("NOTICE TO ALL EMPLOYEES", font_title, 200),
        ("", font_body, 0),
        ("Date: March 20, 2024", font_body, 60),
        ("Re: Updated Safety Protocols - Welding Operations", font_body, 60),
        ("", font_body, 0),
        ("Effective immediately, all welding personnel must comply", font_body, 60),
        ("with the following updated safety requirements:", font_body, 60),
        ("", font_body, 0),
        ("1. Complete the updated AWS safety training module", font_body, 80),
        ("2. Wear approved PPE at all times (face shield, gloves)", font_body, 80),
        ("3. Fire suppression systems must be inspected weekly", font_body, 80),
        ("4. No welding within 35 feet of combustible materials", font_body, 80),
        ("", font_body, 0),
        ("Failure to comply will result in immediate suspension.", font_body, 60),
        ("", font_body, 0),
        ("- Management, BigThorium Energy Corp", font_body, 60),
    ]

    y = 30
    for text, font, x in lines:
        if text:
            draw.text((x, y), text, fill="black", font=font)
        y += 30

    img.save(str(path))
    print(f"  Created: {path.name}")


# ── Main ─────────────────────────────────────────────────────────────
def main():
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    print("=" * 60)
    print(f"Creating sample files in: {SAMPLE_DIR}")
    print("=" * 60)

    create_sample_txt()
    create_sample_pdf()
    create_sample_docx()
    create_sample_eml()
    create_sample_image()

    files = list(SAMPLE_DIR.iterdir())
    print(f"\n{'=' * 60}")
    print(f"Done! {len(files)} files created:")
    for f in sorted(files):
        size = f.stat().st_size
        print(f"  {f.name:40s} ({size:,} bytes)")

    print(f"\nTo ingest:")
    print(f'  python -m app.ingest --dir "{SAMPLE_DIR}" --doc-type "test-docs" --entity-name "BigThorium"')
    print("=" * 60)


if __name__ == "__main__":
    main()
